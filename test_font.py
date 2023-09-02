import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def find_string_in_docx(docx_file, target_string):
    doc = Document(docx_file)
    occurrences = []

    for para in doc.paragraphs:
        if target_string in para.text:
            occurrences.append(para.text)

    return occurrences

def set_font_run(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = font_size
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml("<w:rFonts %s w:hint=\"eastAsia\" />" % nsdecls("w"))
    rPr.insert_element_before(rFonts, "w:sz")
    return run

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script_name.py <target_string>")
        sys.exit(1)

    docx_file_path = "docs/tupi-dic.docx"
    target_string = sys.argv[1]

    found_occurrences = find_string_in_docx(docx_file_path, target_string)
    
    if found_occurrences:
        print("Occurrences found:")
        for idx, occurrence in enumerate(found_occurrences, start=1):
            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run(occurrence)
            run = set_font_run(run, "Arial", Pt(12))  # Use Arial or another font that supports the character
            print(f"{idx}. Paragraph:")
            doc.save("temp.docx")  # Save the paragraph with corrected font
            with open("temp.docx", "rb") as f:
                corrected_paragraph = f.read()
                print(corrected_paragraph.decode("utf-8"))
            print("=" * 20)
    else:
        print("No occurrences found.")
